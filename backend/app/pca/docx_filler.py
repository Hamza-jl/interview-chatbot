"""Write collected answers back into the original Word templates.

The template file is never regenerated - it is opened, filled and saved, so
styles, numbering, headers and the table of contents survive untouched.

Two write modes:

``cell``  a single answer cell (Question / Reponse tables, fiche de suivi).
``rows``  a repeating grid: the blank template rows are consumed first, extra
          rows are cloned from a blank one, and leftover blanks are removed.
"""
from __future__ import annotations

import copy
import datetime as dt
import io
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from app.core.config import settings
from app.pca.blueprint import AUTO_FICHE_ROWS, Question, Target, get_plan

_NB_MARKER = re.compile(r"^\s*N\.?\s*B\.?\s*:", re.IGNORECASE)


# --------------------------------------------------------------------------- #
# Low-level cell helpers
# --------------------------------------------------------------------------- #
def _reference_rpr(cell: _Cell):
    """Formatting of the first real run in the cell, so inserts look native."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run._r.find(qn("w:rPr")) is not None:
                return copy.deepcopy(run._r.find(qn("w:rPr")))
    return None


def _clear_cell(cell: _Cell) -> Paragraph:
    """Drop every paragraph but the first, and empty that one. Returns it."""
    for para in list(cell.paragraphs)[1:]:
        para._p.getparent().remove(para._p)
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._r.getparent().remove(run._r)
    return first


def _write_lines(cell: _Cell, lines: List[str], rpr=None) -> None:
    """Render `lines` into the cell, one paragraph per line."""
    rpr = rpr if rpr is not None else _reference_rpr(cell)
    para = _clear_cell(cell)
    base_p = para._p

    for i, line in enumerate(lines or [""]):
        target = para if i == 0 else Paragraph(copy.deepcopy(base_p), cell)
        if i > 0:
            for run in list(target.runs):
                run._r.getparent().remove(run._r)
            base_p.addnext(target._p)
            base_p = target._p
        run = target.add_run(line)
        if rpr is not None:
            run._r.insert(0, copy.deepcopy(rpr))


def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def set_cell(cell: _Cell, text: str, preserve_notes: bool = True) -> None:
    """Write an answer, keeping any pre-printed 'N.B. :' instruction below it."""
    kept = [
        line.strip()
        for line in _cell_text(cell).splitlines()
        if preserve_notes and _NB_MARKER.match(line)
    ]
    lines = [ln.rstrip() for ln in (text or "").splitlines() if ln.strip()] or [""]
    _write_lines(cell, lines + kept)


# --------------------------------------------------------------------------- #
# Grid helpers
# --------------------------------------------------------------------------- #
def _row_is_blank(row: _Row) -> bool:
    return all(not _cell_text(c) for c in row.cells)


def _clone_row(table: Table, template_row: _Row) -> _Row:
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    row = _Row(new_tr, table)
    for cell in row.cells:
        _clear_cell(cell)
    return row


def fill_grid(table: Table, columns: List[str], rows: List[Dict[str, Any]], header_rows: int = 1) -> None:
    """Fill a repeating table, growing or shrinking it to match the data."""
    body = list(table.rows)[header_rows:]
    if not body:  # degenerate template - nothing to clone from
        return

    blank_template = body[0]
    n_cols = min(len(columns), len(table.columns))

    for i, data in enumerate(rows):
        row = body[i] if i < len(body) else _clone_row(table, blank_template)
        for c in range(n_cols):
            value = data.get(columns[c], "")
            set_cell(row.cells[c], str(value or "").strip(), preserve_notes=False)

    # Remove template rows that stayed empty.
    for leftover in body[len(rows):]:
        if _row_is_blank(leftover):
            leftover._tr.getparent().remove(leftover._tr)


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def _set_entity_heading(doc: Document, structure_name: str) -> None:
    """The cover line reads 'Entite : xxx' in the generic template."""
    for para in doc.paragraphs[:12]:
        text = para.text.strip()
        if text.lower().startswith("entit") and (":" in text):
            prefix, _, current = text.partition(":")
            if current.strip().lower() in {"xxx", ""}:
                rpr = None
                for run in para.runs:
                    if run._r.find(qn("w:rPr")) is not None:
                        rpr = copy.deepcopy(run._r.find(qn("w:rPr")))
                        break
                for run in list(para.runs):
                    run._r.getparent().remove(run._r)
                run = para.add_run(f"{prefix.strip()} : {structure_name}")
                if rpr is not None:
                    run._r.insert(0, rpr)
            return


def fill_document(
    template_path: str,
    template_kind: str,
    answers: Dict[str, Dict[str, Any]],
    *,
    structure_name: str,
    structure_code: str,
    redacteur: str,
    interview_date: Optional[dt.date] = None,
    version: str = "V1.0",
) -> bytes:
    """Return the filled .docx as bytes. `answers` maps question_id -> payload."""
    doc = Document(template_path)
    tables = doc.tables
    plan = get_plan(template_kind)

    def table_at(index_1based: int) -> Optional[Table]:
        idx = index_1based - 1
        return tables[idx] if 0 <= idx < len(tables) else None

    # --- fiche de suivi: the fields the platform knows without asking -------
    fiche = table_at(1)
    if fiche is not None:
        day = interview_date or dt.date.today()
        auto = {
            AUTO_FICHE_ROWS["date"]: day.strftime("%d/%m/%Y"),
            AUTO_FICHE_ROWS["entite"]: structure_name,
            AUTO_FICHE_ROWS["redacteur"]: redacteur,
            AUTO_FICHE_ROWS["version"]: version,
            AUTO_FICHE_ROWS["reference"]: f"{settings.DOC_REFERENCE_PREFIX}-{structure_code}-{version}",
        }
        for row_idx, value in auto.items():
            if row_idx < len(fiche.rows):
                set_cell(fiche.rows[row_idx].cells[1], value)

    _set_entity_heading(doc, structure_name)

    # --- everything the interview collected --------------------------------
    for question in plan:
        payload = answers.get(question.id)
        if not payload:
            continue
        _apply(question, payload, table_at(question.target.table))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply(question: Question, payload: Dict[str, Any], table: Optional[Table]) -> None:
    if table is None:
        return
    target: Target = question.target

    if target.mode == "rows":
        rows = payload.get("rows") or []
        if rows:
            fill_grid(table, [c.id for c in question.columns], rows, target.header_rows)
        return

    value = (payload.get("value") or "").strip()
    if not value:
        return
    if target.row is None or target.row >= len(table.rows):
        return
    row = table.rows[target.row]
    col = min(target.col, len(row.cells) - 1)
    set_cell(row.cells[col], value)
