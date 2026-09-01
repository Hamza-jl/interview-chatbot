"""Google Forms responses -> the client's filled .docx.

The forms are a second collection route onto the same plan and the same
template filler, so the thing worth testing is that a response workbook lands
on the right questions, in the right columns, for the right entity.
"""
from __future__ import annotations

import datetime as dt
import io

import pytest

from app.pca.blueprint import get_plan
from app.scripts import export_forms_spec, from_forms
from app.scripts.seed import STRUCTURES
from tests.conftest import needs_templates


# --------------------------------------------------------------------------- #
# The spec handed to Apps Script
# --------------------------------------------------------------------------- #
def test_the_spec_covers_every_question_and_structure():
    """The forms are generated from the blueprint, never transcribed."""
    spec = export_forms_spec.build_spec()

    assert len(spec["structures"]) == len(STRUCTURES)
    for kind in ("dsi", "entite"):
        plan = get_plan(kind)
        exported = spec["plans"][kind]["questions"]
        assert [q["id"] for q in exported] == [q.id for q in plan], kind
        assert [q["label"] for q in exported] == [q.label for q in plan], kind


def test_every_grid_ships_its_columns():
    """A table question is useless in a form without its column list."""
    spec = export_forms_spec.build_spec()
    for kind, plan in spec["plans"].items():
        for question in plan["questions"]:
            if question["kind"] == "grid":
                assert question["columns"], f"{kind}/{question['id']}"
                assert all(c["label"] for c in question["columns"])


def test_exactly_one_structure_uses_the_long_plan():
    spec = export_forms_spec.build_spec()
    dsi = [s for s in spec["structures"] if s["templateKind"] == "dsi"]
    assert len(dsi) == 1, "the DSI form is the special one"
    assert len(spec["plans"]["dsi"]["questions"]) > len(spec["plans"]["entite"]["questions"])


# --------------------------------------------------------------------------- #
# Matching a sheet back to its entity
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    "tab,expected_code",
    [
        ("État des lieux — Systèmes d'Information", "SI"),
        ("Etat des lieux - Audit Comptable et Financier", "ACF"),
        # Excel caps a sheet name at 31 characters, so a real export arrives
        # cut mid-word. This is what actually comes out of Google.
        ("Etat des lieux - Systèmes d'Inf", "SI"),
        ("Etat des lieux - Audit Comptabl", "ACF"),
        ("Audit Comptable et Financier (réponses)", "ACF"),
    ],
)
def test_a_truncated_tab_name_still_finds_its_entity(tab, expected_code):
    match = from_forms._structure_for(tab)
    assert match is not None, tab
    assert match[0] == expected_code


def test_an_unrelated_tab_matches_nothing():
    assert from_forms._structure_for("Liens") is None
    assert from_forms._structure_for("Feuille 1") is None


# --------------------------------------------------------------------------- #
# Building answers from a response row
# --------------------------------------------------------------------------- #
def _row_for(kind: str, grid_text=None, blank_last: bool = False):
    plan = get_plan(kind)
    header = ["Horodateur", "Adresse e-mail"] + [q.label for q in plan]
    row = [dt.datetime(2026, 8, 31, 9, 30), "a@b.tn"]
    for question in plan:
        if question.kind == "grid":
            row.append(grid_text if grid_text is not None else " | ".join(
                c.choices[0] if c.choices else f"{c.label} A" for c in question.columns
            ))
        else:
            row.append("Mme Leila Ben Salah")
    if blank_last:
        row[-1] = ""
    return plan, header, row


def test_a_response_row_maps_onto_the_plan():
    plan, header, row = _row_for("entite")
    answers, blank, warnings = from_forms.build_answers(plan, header, row)

    assert len(answers) == len(plan)
    assert blank == [] and warnings == []
    grid = next(q for q in plan if q.kind == "grid")
    rows = answers[grid.id]["rows"]
    assert rows and set(rows[0]) == {c.id for c in grid.columns}


def test_a_blank_answer_is_reported_not_invented():
    plan, header, row = _row_for("entite", blank_last=True)
    answers, blank, _ = from_forms.build_answers(plan, header, row)

    assert plan[-1].label in blank
    assert plan[-1].id not in answers, "a blank point must leave the cell empty"


def test_a_table_without_separators_is_kept_and_flagged():
    """Nothing is dropped, but a silent guess is worse than a warning."""
    plan, header, row = _row_for("entite", grid_text="Comptabilite\nTresorerie")
    answers, _, warnings = from_forms.build_answers(plan, header, row)

    grid = next(q for q in plan if q.kind == "grid")
    rows = answers[grid.id]["rows"]
    first = grid.columns[0].id
    assert [r[first] for r in rows] == ["Comptabilite", "Tresorerie"]
    assert all(r[c.id] == "" for r in rows for c in grid.columns[1:])
    assert any("separateur" in from_forms._fold(w) for w in warnings)


def test_a_missing_column_is_reported():
    plan, header, row = _row_for("entite")
    header = header[:-1]                      # the form lost its last question
    row = row[:-1]
    _, _, warnings = from_forms.build_answers(plan, header, row)
    assert any("absente" in w for w in warnings)


# --------------------------------------------------------------------------- #
# End to end
# --------------------------------------------------------------------------- #
@needs_templates
def test_a_workbook_becomes_the_filled_document(tmp_path):
    from openpyxl import Workbook

    import docx

    book = Workbook()
    book.remove(book.active)
    for code, name, _parent, kind in [
        next(s for s in STRUCTURES if s[3] == "dsi"),
        next(s for s in STRUCTURES if s[0] == "ACF"),
    ]:
        plan, header, row = _row_for(kind)
        # Truncated exactly as Excel would truncate it.
        sheet = book.create_sheet(f"Etat des lieux - {name}"[:31])
        sheet.append(header)
        sheet.append(row)

    source = tmp_path / "reponses.xlsx"
    book.save(source)
    out = tmp_path / "sorties"

    assert from_forms.generate(str(source), str(out), "Tests") == 0

    produced = sorted(p.name for p in out.iterdir())
    assert len(produced) == 2, produced

    document = docx.Document(str(out / "Etat_des_lieux_Audit_Comptable_et_Financier.docx"))
    fiche = "\n".join(c.text for row in document.tables[0].rows for c in row.cells)
    assert "Audit Comptable et Financier" in fiche
    assert "Mme Leila Ben Salah" in fiche

    # A grid answer must land as table rows, not as one blob in a cell.
    grid_tables = [t for t in document.tables if len(t.columns) >= 3]
    assert grid_tables, "no table wide enough to be a grid"
    assert any(len(t.rows) > 1 for t in grid_tables)


@needs_templates
def test_the_most_recent_response_wins(tmp_path):
    from openpyxl import Workbook

    import docx

    plan, header, row = _row_for("entite")
    old = list(row)
    old[0] = dt.datetime(2026, 8, 1, 8, 0)
    old[2] = "M. Ancien Responsable"
    new = list(row)
    new[0] = dt.datetime(2026, 8, 30, 8, 0)
    new[2] = "Mme Nouvelle Responsable"

    book = Workbook()
    book.remove(book.active)
    sheet = book.create_sheet("Etat des lieux - Audit Comptabl")
    sheet.append(header)
    sheet.append(old)          # deliberately out of order
    sheet.append(new)

    source = tmp_path / "reponses.xlsx"
    book.save(source)
    out = tmp_path / "sorties"
    assert from_forms.generate(str(source), str(out), "Tests") == 0

    produced = list(out.iterdir())
    assert len(produced) == 1, "one document per entity by default"
    text = "\n".join(
        c.text for t in docx.Document(str(produced[0])).tables for r in t.rows for c in r.cells
    )
    assert "Mme Nouvelle Responsable" in text
    assert "M. Ancien Responsable" not in text


# --------------------------------------------------------------------------- #
# The two-form layout: one entity form, one DSI form, entity chosen in the form
# --------------------------------------------------------------------------- #
STRUCTURE_COLUMN = from_forms.STRUCTURE_COLUMN


def _entity_row(kind: str, code: str, name: str, responsable: str, when: dt.datetime):
    """A response as the single entity form writes it: dropdown, then the plan."""
    plan = get_plan(kind)
    header = ["Horodateur", "Adresse e-mail", STRUCTURE_COLUMN] + [q.label for q in plan]
    row = [when, "a@b.tn", f"{name} ({code})"]
    for question in plan:
        if question.kind == "grid":
            row.append(" | ".join(
                c.choices[0] if c.choices else f"{c.label} A" for c in question.columns
            ))
        elif question.kind == "field":
            row.append(responsable)
        else:
            row.append("Reponse redigee.")
    return header, row


@pytest.mark.parametrize(
    "answer,expected",
    [
        ("Audit Comptable et Financier (ACF)", "ACF"),
        ("Systèmes d'Information (SI)", "SI"),
        ("  Gestion administrative Capital Humain (GCH)  ", "GCH"),
        # The name alone still resolves, if someone edits the choice list.
        ("Audit Comptable et Financier", "ACF"),
    ],
)
def test_the_dropdown_answer_identifies_the_entity(answer, expected):
    match = from_forms.structure_from_answer(answer)
    assert match is not None, answer
    assert match[0] == expected


def test_an_empty_or_unknown_choice_resolves_to_nothing():
    assert from_forms.structure_from_answer("") is None
    assert from_forms.structure_from_answer("Entité inventée (ZZZ)") is None


def test_the_derived_per_entity_tabs_are_named_by_code():
    """Excel truncates to 31 characters, so the code has to lead."""
    assert from_forms._structure_for("ACF - Audit Comptable et Fin")[0] == "ACF"
    assert from_forms._structure_for("SI - Systemes d'Information")[0] == "SI"


def test_one_sheet_of_many_entities_is_split_by_the_dropdown():
    header, acf = _entity_row("entite", "ACF", "Audit Comptable et Financier",
                              "Mme Leila Ben Salah", dt.datetime(2026, 8, 30, 9, 0))
    _, gch = _entity_row("entite", "GCH", "Gestion administrative Capital Humain",
                         "M. Sami Gharbi", dt.datetime(2026, 8, 30, 10, 0))

    buckets, unmatched = from_forms.collect([("Réponses au formulaire 1", header, [acf, gch])])

    assert unmatched == []
    assert set(buckets) == {"ACF", "GCH"}
    assert len(buckets["ACF"]) == 1 and len(buckets["GCH"]) == 1


def test_the_derived_tabs_do_not_duplicate_the_raw_ones():
    """The per-entity tabs repeat the raw rows; one response, one document."""
    header, acf = _entity_row("entite", "ACF", "Audit Comptable et Financier",
                              "Mme Leila Ben Salah", dt.datetime(2026, 8, 30, 9, 0))

    buckets, _ = from_forms.collect([
        ("Réponses au formulaire 1", header, [acf]),      # raw, as Google writes it
        ("ACF - Audit Comptable et Fin", header, [acf]),  # the routed copy
    ])

    assert list(buckets) == ["ACF"]
    assert len(buckets["ACF"]) == 1, "the same response must not count twice"


def test_the_two_forms_land_in_the_same_run():
    header_e, acf = _entity_row("entite", "ACF", "Audit Comptable et Financier",
                                "Mme Leila Ben Salah", dt.datetime(2026, 8, 30, 9, 0))
    header_d, si = _entity_row("dsi", "SI", "Systèmes d'Information",
                               "M. Karim Trabelsi", dt.datetime(2026, 8, 30, 9, 0))

    buckets, _ = from_forms.collect([
        ("Réponses au formulaire 1", header_e, [acf]),
        ("Réponses au formulaire 2", header_d, [si]),
    ])
    assert set(buckets) == {"ACF", "SI"}


@needs_templates
def test_two_forms_produce_one_document_per_entity(tmp_path):
    from openpyxl import Workbook

    import docx

    header_e, acf = _entity_row("entite", "ACF", "Audit Comptable et Financier",
                                "Mme Leila Ben Salah", dt.datetime(2026, 8, 30, 9, 0))
    _, gch_old = _entity_row("entite", "GCH", "Gestion administrative Capital Humain",
                             "M. Ancien", dt.datetime(2026, 8, 1, 9, 0))
    _, gch_new = _entity_row("entite", "GCH", "Gestion administrative Capital Humain",
                             "Mme Nouvelle", dt.datetime(2026, 8, 29, 9, 0))
    header_d, si = _entity_row("dsi", "SI", "Systèmes d'Information",
                               "M. Karim Trabelsi", dt.datetime(2026, 8, 30, 9, 0))

    book = Workbook()
    book.remove(book.active)
    entities = book.create_sheet("Reponses au formulaire 1")
    entities.append(header_e)
    for row in (acf, gch_old, gch_new):
        entities.append(row)
    dsi = book.create_sheet("Reponses au formulaire 2")
    dsi.append(header_d)
    dsi.append(si)
    routed = book.create_sheet("ACF - Audit Comptable et Fin")   # the derived view
    routed.append(header_e)
    routed.append(acf)

    source = tmp_path / "reponses.xlsx"
    book.save(source)
    out = tmp_path / "sorties"
    assert from_forms.generate(str(source), str(out), "Tests") == 0

    produced = sorted(p.name for p in out.iterdir())
    assert produced == [
        "Etat_des_lieux_Audit_Comptable_et_Financier.docx",
        "Etat_des_lieux_Gestion_administrative_Capital_Humain.docx",
        "Etat_des_lieux_Systemes_d_Information.docx",
    ], produced

    def text_of(name):
        document = docx.Document(str(out / name))
        return "\n".join(c.text for t in document.tables for r in t.rows for c in r.cells)

    gch = text_of("Etat_des_lieux_Gestion_administrative_Capital_Humain.docx")
    assert "Mme Nouvelle" in gch and "M. Ancien" not in gch, "most recent wins"

    si_text = text_of("Etat_des_lieux_Systemes_d_Information.docx")
    assert "Systèmes d'Information" in si_text
