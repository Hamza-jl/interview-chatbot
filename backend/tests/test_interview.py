"""The interview itself: routing, extraction, template fidelity and export."""
from __future__ import annotations

import hashlib
import io

import docx

from app.pca.blueprint import ENTITE_PLAN, TEMPLATE_FILES, get_plan, sections
from tests.conftest import (
    API,
    answer,
    authenticate,
    confirm,
    finish,
    needs_templates,
    unaccented,
)


def _open_session(client, headers, code: str = "DSI") -> dict:
    structures = client.get(f"{API}/structures", headers=headers).json()
    target = next(s for s in structures if s["code"] == code)
    return client.post(f"{API}/sessions", headers=headers, json={"structure_id": target["id"]}).json()


def _say(client, headers, session_id: str, message: str) -> dict:
    res = client.post(
        f"{API}/sessions/{session_id}/messages", headers=headers, json={"message": message}
    )
    assert res.status_code == 200, res.text
    return res.json()


# --------------------------------------------------------------------------- #
# Blueprint fidelity - the plan must match the shipped templates exactly
# --------------------------------------------------------------------------- #
@needs_templates
def test_every_question_targets_a_real_table_and_cell():
    import os

    from app.core.config import settings

    for kind, filename in TEMPLATE_FILES.items():
        document = docx.Document(os.path.join(settings.TEMPLATE_DIR, filename))
        tables = document.tables
        for question in get_plan(kind):
            target = question.target
            assert 1 <= target.table <= len(tables), f"{question.id}: table {target.table} missing"
            table = tables[target.table - 1]

            if target.mode == "cell":
                assert target.row is not None and target.row < len(table.rows), question.id
                assert target.col < len(table.rows[target.row].cells), question.id
            else:
                # A grid question must declare exactly as many columns as the table has.
                assert len(question.columns) == len(table.columns), (
                    f"{question.id}: {len(question.columns)} columns declared, "
                    f"table has {len(table.columns)}"
                )


def test_no_two_questions_write_to_the_same_cell():
    for kind in TEMPLATE_FILES:
        seen: set[tuple] = set()
        for question in get_plan(kind):
            slot = (question.target.table, question.target.row, question.target.col,
                    question.target.mode)
            assert slot not in seen, f"{question.id} collides with an earlier question"
            seen.add(slot)


def test_question_ids_are_unique_and_sections_are_ordered():
    for kind in TEMPLATE_FILES:
        plan = get_plan(kind)
        assert len({q.id for q in plan}) == len(plan)
        assert sections(kind)[0] == "Fiche de suivi"


def test_entite_plan_is_shorter_than_the_dsi_plan():
    assert len(ENTITE_PLAN) < len(get_plan("dsi"))


# --------------------------------------------------------------------------- #
# Conversational routing
# --------------------------------------------------------------------------- #
def test_a_definition_question_is_answered_without_recording_anything(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Qu'est-ce que la criticite SI exactement ?")

    assert turn["intent"] == "question"
    assert turn["recorded"] is False
    assert turn["state"]["answered"] == 0
    assert turn["state"]["cursor"] == state["cursor"], "the interview must not advance"
    assert "Vitale" in turn["reply"]["body"], "the answer comes from the template glossary"


def test_an_answer_becomes_a_draft_and_waits_for_confirmation(client):
    """Nothing reaches the document until the interviewee has verified it."""
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Mme Sonia Ben Ammar")

    assert turn["pending"] is not None, "the extraction must be offered for review"
    assert turn["pending"]["question_id"] == "dsi.fiche.responsable"
    assert "Ben Ammar" in turn["pending"]["value"]
    assert turn["recorded"] is False, "a draft is not a recorded answer"
    assert turn["state"]["answered"] == 0, "a draft is not progress"
    assert turn["state"]["cursor"] == state["cursor"], "the interview must hold"


def test_confirming_records_the_answer_and_advances(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Mme Sonia Ben Ammar")
    done = confirm(client, headers, state["id"], turn["pending"])

    assert done["recorded"] is True
    assert done["state"]["answered"] == 1
    assert done["state"]["cursor"] == state["cursor"] + 1


def test_the_panel_edit_is_what_gets_stored(client):
    """The interviewee overrules the engine - that is the point of the step."""
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Mme Sonia Ben Ammar")
    confirm(client, headers, state["id"], turn["pending"], value="Mme Sonia BEN AMMAR, DSI")

    answers = client.get(f"{API}/sessions/{state['id']}/answers", headers=headers).json()
    stored = next(a for a in answers if a["question_id"] == "dsi.fiche.responsable")
    assert stored["value"] == "Mme Sonia BEN AMMAR, DSI"
    assert stored["confirmed"] is True


@needs_templates
def test_grid_edits_survive_confirmation_and_drop_unknown_columns(client):
    headers = authenticate(client, "client@example.com")
    detail = _open_session(client, headers)
    session_id = detail["state"]["id"]
    grid = next(q for q in get_plan("dsi") if q.kind == "grid")

    turn = None
    for _ in range(30):
        state = turn["state"] if turn else detail["state"]
        if state["question"] and state["question"]["id"] == grid.id:
            break
        turn = answer(client, headers, session_id, "Reponse redigee pour ce point.")

    draft = _say(client, headers, session_id, grid.example)
    assert draft["pending"]["kind"] == "grid"

    edited = [{grid.columns[0].id: "Corrige", "colonne_inventee": "ignore"}]
    done = confirm(client, headers, session_id, draft["pending"], rows=edited)
    assert done["recorded"] is True

    answers = client.get(f"{API}/sessions/{session_id}/answers", headers=headers).json()
    stored = next(a for a in answers if a["question_id"] == grid.id)
    assert stored["rows"][0][grid.columns[0].id] == "Corrige"
    assert "colonne_inventee" not in stored["rows"][0], "column names are fixed by the template"


def test_a_discarded_draft_leaves_nothing_behind(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Mme Sonia Ben Ammar")
    res = client.post(
        f"{API}/sessions/{state['id']}/discard",
        headers=headers,
        json={"question_id": turn["pending"]["question_id"]},
    )
    assert res.status_code == 200
    assert res.json()["state"]["cursor"] == state["cursor"], "stays on the question"

    answers = client.get(f"{API}/sessions/{state['id']}/answers", headers=headers).json()
    stored = next(a for a in answers if a["question_id"] == "dsi.fiche.responsable")
    assert stored["completeness"] == "vide"
    assert stored["confirmed"] is False


def test_an_empty_confirmation_is_refused(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]
    _say(client, headers, state["id"], "Mme Sonia Ben Ammar")

    res = client.post(
        f"{API}/sessions/{state['id']}/confirm",
        headers=headers,
        json={"question_id": "dsi.fiche.responsable", "value": "   "},
    )
    assert res.status_code == 422


@needs_templates
def test_an_unconfirmed_draft_never_reaches_the_document(client):
    """The whole point of the step: a draft must not appear in the deliverable."""
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    answer(client, headers, state["id"], "Mme Sonia Ben Ammar")   # confirmed
    _say(client, headers, state["id"], "M. Karim Trabelsi")        # left as a draft

    finish(client, headers, state["id"])                           # closed with gaps
    export = client.post(f"{API}/sessions/{state['id']}/export", headers=headers)
    assert export.status_code == 200
    document = client.get(
        f"{API}/exports/{export.json()['download_token']}", headers=headers
    ).content

    import docx as _docx

    fiche = _docx.Document(io.BytesIO(document)).tables[0]
    text = "\n".join(cell.text for row in fiche.rows for cell in row.cells)
    assert "Ben Ammar" in text, "the confirmed answer is there"
    assert "Trabelsi" not in text, "the unconfirmed draft is not"


def test_an_unknown_term_is_not_invented(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]

    turn = _say(client, headers, state["id"], "Que signifie le terme zorglub dans ce contexte ?")

    assert turn["recorded"] is False
    assert "referentiel" in unaccented(turn["reply"]["body"])


def test_the_interview_never_loops_forever_on_one_question(client):
    """Unusable input for a table is refused, then the interview moves on."""
    headers = authenticate(client, "client@example.com")
    detail = _open_session(client, headers)
    session_id = detail["state"]["id"]

    # walk to the first grid question
    turn = None
    for _ in range(30):
        state = turn["state"] if turn else detail["state"]
        if state["question"] and state["question"]["kind"] == "grid":
            break
        turn = answer(client, headers, session_id, "Reponse redigee pour ce point.")

    cursor_before = turn["state"]["cursor"]
    for _ in range(4):
        turn = answer(client, headers, session_id, "je ne sais pas trop quoi mettre ici")

    assert turn["state"]["cursor"] > cursor_before, "must not stall on the same question"


# --------------------------------------------------------------------------- #
# The review gate - nothing closes with silent holes in it
# --------------------------------------------------------------------------- #
def _fill(client, headers, session_id: str, answer_row: dict) -> None:
    """Answer one point straight from the review panel, whatever its kind."""
    body = {"question_id": answer_row["question_id"]}
    if answer_row["kind"] == "grid":
        body["rows"] = [{c["id"]: "a renseigner" for c in answer_row["columns"]}]
    else:
        body["value"] = "Valeur fournie pendant la revue finale."
    res = client.put(f"{API}/sessions/{session_id}/answers", headers=headers, json=body)
    assert res.status_code == 200, res.text


def test_leaving_questions_blank_holds_the_interview_open_for_review(client):
    """Running out of questions is not the same thing as having answered them."""
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]

    answer(client, headers, session_id, "Mme Sonia Ben Ammar")
    turn = _say(client, headers, session_id, "je veux terminer maintenant")

    assert turn["completed"] is False, "the interview must not close over blank points"
    state = turn["state"]
    assert state["status"] == "in_progress"
    assert state["awaiting_review"] is True
    assert len(state["missing"]) == state["total"] - state["answered"] > 0
    # The recap has to name them, not merely count them.
    assert state["missing"][0]["label"] in turn["reply"]["body"]


def test_a_blank_point_carries_its_columns_so_it_can_be_answered(client):
    """The point most likely to be opened is the one with nothing stored yet."""
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]

    rows = client.get(f"{API}/sessions/{session_id}/answers", headers=headers).json()
    for row in rows:
        if row["kind"] == "grid":
            assert row["rows"] in (None, []), "nothing answered yet"
            assert row["columns"], f"{row['question_id']}: no columns to render"
            assert row["prompt"], f"{row['question_id']}: no prompt to show"
            break
    else:
        raise AssertionError("the plan has no grid question")


def test_the_document_cannot_be_produced_before_the_interview_is_closed(client):
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]
    answer(client, headers, session_id, "Mme Sonia Ben Ammar")

    refused = client.post(f"{API}/sessions/{session_id}/export", headers=headers)
    assert refused.status_code == 409

    finish(client, headers, session_id)
    assert client.post(f"{API}/sessions/{session_id}/export", headers=headers).status_code == 200


def test_closing_over_blank_points_takes_an_explicit_acknowledgement(client):
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]
    answer(client, headers, session_id, "Mme Sonia Ben Ammar")

    refused = finish(client, headers, session_id, acknowledge=False, expect=409)
    assert "sans reponse" in unaccented(refused["detail"])

    closed = finish(client, headers, session_id, acknowledge=True)
    assert closed["completed"] is True
    assert closed["state"]["status"] == "completed"
    assert closed["state"]["awaiting_review"] is False


def test_answering_from_the_review_clears_the_gate(client):
    """Filling the last hole lets the interview close without an override."""
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]
    _say(client, headers, session_id, "je veux terminer maintenant")

    rows = client.get(f"{API}/sessions/{session_id}/answers", headers=headers).json()
    for row in rows:
        _fill(client, headers, session_id, row)

    state = client.get(f"{API}/sessions/{session_id}", headers=headers).json()["state"]
    assert state["missing"] == []
    assert state["awaiting_review"] is True, "still open until it is deliberately closed"

    closed = finish(client, headers, session_id, acknowledge=False)
    assert closed["completed"] is True


def test_the_composer_is_closed_while_the_review_is_pending(client):
    """A message sent past the last question must not be filed against it."""
    headers = authenticate(client, "client@example.com")
    session_id = _open_session(client, headers)["state"]["id"]
    _say(client, headers, session_id, "je veux terminer maintenant")

    res = client.post(
        f"{API}/sessions/{session_id}/messages", headers=headers,
        json={"message": "encore une remarque"},
    )
    assert res.status_code == 409


# --------------------------------------------------------------------------- #
# One interview per entity
# --------------------------------------------------------------------------- #
def _start(client, headers, code: str = "DSI"):
    structures = client.get(f"{API}/structures", headers=headers).json()
    target = next(s for s in structures if s["code"] == code)
    return client.post(f"{API}/sessions", headers=headers, json={"structure_id": target["id"]})


def test_an_interview_left_open_is_resumed_not_duplicated(client):
    headers = authenticate(client, "client@example.com")
    first = _start(client, headers).json()["state"]
    answer(client, headers, first["id"], "Mme Sonia Ben Ammar")

    again = _start(client, headers).json()["state"]

    assert again["id"] == first["id"], "picking the structure again must resume"
    assert again["answered"] == 1, "and pick up where it was left off"
    assert again["cursor"] > 0


def test_a_closed_interview_cannot_be_restarted(client):
    """The document is the deliverable - a second run would compete with it."""
    headers = authenticate(client, "client@example.com")
    session_id = _start(client, headers).json()["state"]["id"]
    answer(client, headers, session_id, "Mme Sonia Ben Ammar")
    finish(client, headers, session_id)

    refused = _start(client, headers)
    assert refused.status_code == 409
    assert "cloture" in unaccented(refused.json()["detail"]).lower()

    # and no second session was conjured for that structure
    listed = client.get(f"{API}/sessions", headers=headers).json()
    for_structure = [s for s in listed if s["structure"]["code"] == "DSI"]
    assert len(for_structure) == 1
    assert for_structure[0]["id"] == session_id


def test_a_closed_interview_is_still_readable_and_downloadable(client):
    headers = authenticate(client, "client@example.com")
    session_id = _start(client, headers).json()["state"]["id"]
    answer(client, headers, session_id, "Mme Sonia Ben Ammar")
    finish(client, headers, session_id)

    detail = client.get(f"{API}/sessions/{session_id}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["state"]["completed_at"], "the closure date is reported"

    export = client.post(f"{API}/sessions/{session_id}/export", headers=headers)
    assert export.status_code == 200
    served = client.get(
        f"{API}/exports/{export.json()['download_token']}", headers=headers
    )
    assert served.status_code == 200


def test_manual_override_bypasses_the_model(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]
    question_id = state["question"]["id"]

    res = client.put(
        f"{API}/sessions/{state['id']}/answers",
        headers=headers,
        json={"question_id": question_id, "value": "Valeur saisie manuellement."},
    )
    assert res.status_code == 200
    assert res.json()["value"] == "Valeur saisie manuellement."

    answers = client.get(f"{API}/sessions/{state['id']}/answers", headers=headers).json()
    stored = next(a for a in answers if a["question_id"] == question_id)
    assert stored["value"] == "Valeur saisie manuellement."


def test_grid_override_drops_unknown_columns(client):
    headers = authenticate(client, "client@example.com")
    state = _open_session(client, headers)["state"]
    grid = next(q for q in get_plan("dsi") if q.kind == "grid")

    res = client.put(
        f"{API}/sessions/{state['id']}/answers",
        headers=headers,
        json={
            "question_id": grid.id,
            "rows": [{grid.columns[0].id: "Production", "injected_column": "malicious"}],
        },
    )
    assert res.status_code == 200
    assert "injected_column" not in res.json()["rows"][0]


# --------------------------------------------------------------------------- #
# Export
# --------------------------------------------------------------------------- #
@needs_templates
def test_full_interview_produces_a_faithful_document(client):
    headers = authenticate(client, "client@example.com")
    detail = _open_session(client, headers)
    session_id = detail["state"]["id"]
    plan = get_plan("dsi")

    turn = None
    for _ in range(len(plan) + 5):
        state = turn["state"] if turn else detail["state"]
        question = state["question"]
        if question is None:
            break
        if question["kind"] == "grid":
            message = " | ".join(f"{c['label'][:12]}-1" for c in question["columns"])
        else:
            message = f"Reponse d'atelier pour {question['label']}."
        turn = answer(client, headers, session_id, message)

    assert turn["completed"] is True

    export = client.post(f"{API}/sessions/{session_id}/export", headers=headers)
    assert export.status_code == 200, export.text
    payload = export.json()

    download = client.get(f"{API}/exports/{payload['download_token']}", headers=headers)
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert hashlib.sha256(download.content).hexdigest() == payload["sha256"]

    document = docx.Document(io.BytesIO(download.content))
    assert len(document.tables) == 17

    fiche = document.tables[0]
    assert fiche.rows[1].cells[1].text.strip() == "Direction des Systemes d'Information"
    assert fiche.rows[5].cells[1].text.strip() == "V1.0"
    from app.core.config import settings

    assert fiche.rows[6].cells[1].text.strip() == (
        f"{settings.DOC_REFERENCE_PREFIX}-DSI-V1.0"
    )

    for index, table in enumerate(document.tables, start=1):
        body = table.rows[1:]
        assert any(
            any(cell.text.strip() for cell in row.cells) for row in body
        ), f"table {index} came out empty"


def test_export_refuses_an_interview_with_no_answers(client):
    headers = authenticate(client, "staff@example.org")
    structures = client.get(f"{API}/structures", headers=headers).json()
    target = next(s for s in structures if s["code"] == "DRH")
    session_id = client.post(
        f"{API}/sessions", headers=headers, json={"structure_id": target["id"]}
    ).json()["state"]["id"]

    res = client.post(f"{API}/sessions/{session_id}/export", headers=headers)
    assert res.status_code == 409


@needs_templates
def test_preprinted_template_notes_survive_the_fill(client):
    """Cells carrying a printed 'N.B. :' instruction keep it below the answer."""
    import os

    from app.core.config import settings
    from app.pca.docx_filler import fill_document

    data = fill_document(
        os.path.join(settings.TEMPLATE_DIR, TEMPLATE_FILES["dsi"]),
        "dsi",
        {"dsi.donnees.documents": {"value": "Contrats editeurs dans la GED juridique."}},
        structure_name="Direction des Systemes d'Information",
        structure_code="DSI",
        redacteur="Consultant Devoteam",
    )
    document = docx.Document(io.BytesIO(data))
    cell = document.tables[7].rows[2].cells[1].text
    assert "Contrats editeurs" in cell
    assert "N.B" in cell, "the printed instruction must not be destroyed"
